# utils/annotate_utils.py
import re
from pathlib import Path
from typing import Dict, List, Optional
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.enum.text import WD_COLOR_INDEX

def _mk_ws_rx(s: str) -> re.Pattern:
    # whitespace-tolerant regex for matching across line wraps
    patt = re.escape(s.strip())
    patt = re.sub(r"\\\s+", r"\\s+", patt)
    return re.compile(patt, flags=re.IGNORECASE)

def annotate_pdf_spans(original_path: str, spans: List[dict], out_dir: Path) -> Optional[str]:
    try:
        doc = fitz.open(original_path)
        for sp in spans:
            text = sp.get("text","")
            if not text or len(text) < 3:
                continue
            page = sp.get("page")
            pages = [doc[page]] if isinstance(page, int) and 0 <= page < len(doc) else list(doc)
            # try exact phrase; if too long, also try halves
            cands = [text]
            words = text.split()
            if len(words) > 20:
                mid = len(words)//2
                cands += [" ".join(words[:mid]), " ".join(words[mid:])]
            for pg in pages:
                for cand in cands:
                    hits = pg.search_for(cand, quads=True)
                    for h in hits:
                        pg.add_highlight_annot(h)
        out_dir.mkdir(parents=True, exist_ok=True)
        outp = out_dir / Path(original_path).name
        doc.save(outp, incremental=False, deflate=True)
        doc.close()
        return str(outp)
    except Exception as e:
        print(f"[WARN] PDF annotate failed for {original_path}: {e}")
        return None

def annotate_docx_spans(original_path: str, spans: List[dict], out_dir: Path) -> Optional[str]:
    try:
        doc = DocxDocument(original_path)
        patterns = [_mk_ws_rx(sp["text"]) for sp in spans if sp.get("text")]
        if not patterns:
            outp = out_dir / Path(original_path).name
            out_dir.mkdir(parents=True, exist_ok=True)
            doc.save(outp)
            return str(outp)
        needles = re.compile("|".join(p.pattern for p in patterns), flags=re.IGNORECASE)
        for para in doc.paragraphs:
            txt = para.text
            if not txt or not needles.search(txt):
                continue
            para_text = txt
            para.clear()
            pos = 0
            for m in needles.finditer(para_text):
                pre = para_text[pos:m.start()]
                hit = para_text[m.start():m.end()]
                if pre: para.add_run(pre)
                r = para.add_run(hit); r.font.highlight_color = WD_COLOR_INDEX.YELLOW
                pos = m.end()
            tail = para_text[pos:]
            if tail: para.add_run(tail)
        out_dir.mkdir(parents=True, exist_ok=True)
        outp = out_dir / Path(original_path).name
        doc.save(outp)
        return str(outp)
    except Exception as e:
        print(f"[WARN] DOCX annotate failed for {original_path}: {e}")
        return None

def annotate_text_sidecar_spans(original_path: str, spans: List[dict], out_dir: Path) -> Optional[str]:
    try:
        text = Path(original_path).read_text(encoding="utf-8", errors="ignore")
        for sp in spans:
            s = sp.get("text","").strip()
            if not s: continue
            patt = _mk_ws_rx(s)
            text = patt.sub(lambda m: f"<<EVIDENCE_START>>{m.group(0)}<<EVIDENCE_END>>", text)
        out_dir.mkdir(parents=True, exist_ok=True)
        outp = out_dir / (Path(original_path).stem + ".evidence.txt")
        outp.write_text(text, encoding="utf-8")
        return str(outp)
    except Exception as e:
        print(f"[WARN] Text annotate failed for {original_path}: {e}")
        return None

def annotate_sources_spans(evidence_map: Dict[str, dict], out_dir: Path) -> Dict[str, dict]:
    """
    evidence_map: { source_path: { 'spans': [ {'text':..., 'page':..., 'chunk_id':...}, ... ] } }
    Returns { source_path: { 'annotated_copy': str|None, 'spans': [...] } }
    """
    result = {}
    for src, v in evidence_map.items():
        spans = v.get("spans", [])
        lower = src.lower()
        if lower.endswith(".pdf"):
            outp = annotate_pdf_spans(src, spans, out_dir)
        elif lower.endswith(".docx"):
            outp = annotate_docx_spans(src, spans, out_dir)
        else:
            outp = annotate_text_sidecar_spans(src, spans, out_dir)
        result[src] = {"annotated_copy": outp, "spans": spans}
    return result
